"""
Export Word (DOCX) — rapport structuré des corrections éditoriales.

Génère un document Word lisible et imprimable à partir des corrections
stockées en base de données pour un job donné.

Structure du document :
  1. Page de titre  : titre, nom du fichier, date, type de document, stats globales
  2. Tableau de synthèse : Catégorie | Nb corrections | Description
  3. Sections par catégorie (A → H) : titre de section coloré + tableau des corrections
     Colonnes : Page | Texte original | Correction proposée | Confiance | Explication

Dépendance : python-docx  (pip install python-docx)
"""
from __future__ import annotations

import logging
from datetime import date
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

from backend.config import settings

logger = logging.getLogger(__name__)

# ── Configuration des catégories ───────────────────────────────────────────────
# Couleurs en RGB (0-255) pour les titres de section Word.
# Correspondance visuelle avec les annotations PDF du pipeline.

_CATEGORY_META: dict[str, dict[str, Any]] = {
    "A": {
        "name": "A – Orthographe",
        "rgb": RGBColor(0xCC, 0x00, 0x00),       # rouge
        "description": "Erreurs d'orthographe lexicale",
    },
    "B": {
        "name": "B – Grammaire",
        "rgb": RGBColor(0xE6, 0x6B, 0x00),       # orange
        "description": "Erreurs grammaticales (accords, conjugaison…)",
    },
    "C": {
        "name": "C – Typographie",
        "rgb": RGBColor(0x66, 0x00, 0xCC),       # violet
        "description": "Problèmes typographiques (espaces, ponctuation, guillemets…)",
    },
    "D": {
        "name": "D – Syntaxe",
        "rgb": RGBColor(0x00, 0x4C, 0xCC),       # bleu
        "description": "Erreurs ou lourdeurs syntaxiques",
    },
    "E": {
        "name": "E – Sémantique",
        "rgb": RGBColor(0x00, 0x80, 0x1A),       # vert
        "description": "Problèmes sémantiques (contresens, répétitions, registre…)",
    },
    "F": {
        "name": "F – Uniformisation",
        "rgb": RGBColor(0x00, 0x80, 0x99),       # cyan
        "description": "Incohérences dans les choix typographiques ou orthographiques",
    },
    "G": {
        "name": "G – Renvois",
        "rgb": RGBColor(0xCC, 0x00, 0x66),       # rose
        "description": "Références internes, notes et renvois",
    },
    "H": {
        "name": "H – Vérification des faits",
        "rgb": RGBColor(0xB8, 0x86, 0x00),       # jaune foncé (lisible sur blanc)
        "description": "Anomalies factuelles (noms propres, dates, titres d'œuvres)",
    },
}

_CATEGORY_ORDER = list("ABCDEFGH")
_BODY_FONT = "Calibri"
_BODY_SIZE = Pt(11)
_SECTION_TITLE_SIZE = Pt(14)
_EXPLANATION_MAX_CHARS = 200

# ── Utilitaires bas niveau ─────────────────────────────────────────────────────


def _set_cell_bg(cell, hex_color: str) -> None:
    """Applique une couleur de fond à une cellule de tableau (XML direct)."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _set_cell_border(cell, border_color: str = "CCCCCC") -> None:
    """Ajoute une bordure fine sur toutes les faces d'une cellule."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), border_color)
        tcBorders.append(el)
    tcPr.append(tcBorders)


def _paragraph_spacing(paragraph, before: int = 0, after: int = 0) -> None:
    """Définit l'espacement avant/après un paragraphe (en points)."""
    pPr = paragraph._p.get_or_add_pPr()
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:before"), str(before * 20))  # twips (1 pt = 20 twips)
    spacing.set(qn("w:after"), str(after * 20))
    pPr.append(spacing)


def _add_page_break(doc: Document) -> None:
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    run.add_break(docx_break_type="page")  # type: ignore[call-arg]
    # Fallback XML si l'API add_break ne supporte pas le type
    from docx.oxml.ns import qn as _qn
    from docx.oxml import OxmlElement as _OxmlElement
    br = _OxmlElement("w:br")
    br.set(_qn("w:type"), "page")
    run._r.append(br)


# ── Construction du document ───────────────────────────────────────────────────


def _add_title_page(
    doc: Document,
    job: object,
    corrections: list,
    cat_counts: dict[str, int],
) -> None:
    """Ajoute la page de titre avec les statistiques globales."""
    doc.add_paragraph()  # espace initial

    # Titre principal
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _paragraph_spacing(title_para, before=60, after=12)
    run = title_para.add_run("Rapport de correction éditoriale")
    run.bold = True
    run.font.name = _BODY_FONT
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(0x1F, 0x39, 0x7A)

    # Nom du fichier
    filename_para = doc.add_paragraph()
    filename_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _paragraph_spacing(filename_para, after=4)
    run = filename_para.add_run(getattr(job, "filename", "Document inconnu"))
    run.bold = True
    run.font.name = _BODY_FONT
    run.font.size = Pt(14)

    # Date de génération
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _paragraph_spacing(date_para, after=4)
    run = date_para.add_run(f"Généré le {date.today().strftime('%d/%m/%Y')}")
    run.font.name = _BODY_FONT
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    # Type de document
    doc_type_labels = {
        "roman": "Roman / texte littéraire",
        "manuel_scolaire": "Manuel scolaire",
        "essai": "Essai / rapport",
        "autre": "Document général",
    }
    doc_type_raw = getattr(job, "doc_type", "autre") or "autre"
    doc_type_label = doc_type_labels.get(doc_type_raw, doc_type_raw.capitalize())

    type_para = doc.add_paragraph()
    type_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _paragraph_spacing(type_para, after=32)
    run = type_para.add_run(f"Type de document : {doc_type_label}")
    run.italic = True
    run.font.name = _BODY_FONT
    run.font.size = Pt(11)

    # Ligne de séparation visuelle
    sep = doc.add_paragraph("─" * 55)
    sep.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _paragraph_spacing(sep, after=16)

    # Stats globales
    total = len(corrections)
    pages_count = getattr(job, "pages_count", None)
    word_count = getattr(job, "word_count", None)

    stats_lines = [
        f"Nombre total de corrections : {total}",
    ]
    if pages_count:
        stats_lines.append(f"Pages analysées : {pages_count}")
    if word_count:
        stats_lines.append(f"Mots analysés : {word_count:,}")

    for line in stats_lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _paragraph_spacing(p, after=4)
        run = p.add_run(line)
        run.font.name = _BODY_FONT
        run.font.size = _BODY_SIZE

    # Répartition par catégorie (liste condensée)
    if cat_counts:
        doc.add_paragraph()
        breakdown_para = doc.add_paragraph()
        breakdown_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = breakdown_para.add_run("Répartition par catégorie")
        run.bold = True
        run.font.name = _BODY_FONT
        run.font.size = _BODY_SIZE
        _paragraph_spacing(breakdown_para, before=12, after=8)

        for cat in _CATEGORY_ORDER:
            count = cat_counts.get(cat, 0)
            if count == 0:
                continue
            meta = _CATEGORY_META.get(cat, {"name": cat, "rgb": RGBColor(0, 0, 0)})
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _paragraph_spacing(p, after=2)
            run = p.add_run(f"{meta['name']} : {count}")
            run.font.name = _BODY_FONT
            run.font.size = _BODY_SIZE
            run.font.color.rgb = meta["rgb"]


def _add_summary_table(
    doc: Document,
    cat_counts: dict[str, int],
) -> None:
    """Ajoute le tableau de synthèse Catégorie | Nb corrections | Description."""
    heading = doc.add_heading("Tableau de synthèse", level=1)
    _paragraph_spacing(heading, before=20, after=10)

    # Entêtes
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    headers = ["Catégorie", "Corrections", "Description"]

    for i, hdr in enumerate(headers):
        hdr_cells[i].text = hdr
        _set_cell_bg(hdr_cells[i], "1F397A")
        para = hdr_cells[i].paragraphs[0]
        run = para.runs[0] if para.runs else para.add_run(hdr)
        run.font.name = _BODY_FONT
        run.font.size = _BODY_SIZE
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    # Données
    for cat in _CATEGORY_ORDER:
        count = cat_counts.get(cat, 0)
        if count == 0:
            continue
        meta = _CATEGORY_META.get(cat, {"name": cat, "description": "", "rgb": RGBColor(0, 0, 0)})
        row_cells = table.add_row().cells
        row_cells[0].text = meta["name"]
        row_cells[1].text = str(count)
        row_cells[2].text = meta.get("description", "")

        for cell in row_cells:
            _set_cell_border(cell)
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.name = _BODY_FONT
                    run.font.size = _BODY_SIZE

        # Couleur de catégorie sur la première cellule
        for run in row_cells[0].paragraphs[0].runs:
            run.bold = True
            run.font.color.rgb = meta["rgb"]

    doc.add_paragraph()


def _add_category_section(
    doc: Document,
    cat: str,
    cat_corrections: list,
) -> None:
    """Ajoute une section pour une catégorie : titre coloré + tableau des corrections."""
    meta = _CATEGORY_META.get(cat, {"name": cat, "rgb": RGBColor(0x33, 0x33, 0x33)})

    # Titre de section
    heading_para = doc.add_paragraph()
    _paragraph_spacing(heading_para, before=20, after=8)
    run = heading_para.add_run(meta["name"])
    run.bold = True
    run.font.name = _BODY_FONT
    run.font.size = _SECTION_TITLE_SIZE
    run.font.color.rgb = meta["rgb"]

    # Note spéciale pour catégorie H
    if cat == "H":
        note_para = doc.add_paragraph()
        _paragraph_spacing(note_para, after=6)
        note_run = note_para.add_run(
            "Note : certaines corrections H peuvent ne pas être visibles dans le PDF annoté "
            "si le texte source est dans une zone rasterisée (image). "
            "Ces corrections sont néanmoins listées ci-dessous et doivent être vérifiées manuellement. "
            "La colonne « Source web » indique la référence consultée lors de la vérification des faits."
        )
        note_run.italic = True
        note_run.font.name = _BODY_FONT
        note_run.font.size = Pt(9)
        note_run.font.color.rgb = RGBColor(0x88, 0x55, 0x00)

    # Tableau des corrections
    is_h = cat == "H"
    num_cols = 6 if is_h else 5
    table = doc.add_table(rows=1, cols=num_cols)
    table.style = "Table Grid"

    # En-tête du tableau
    hdr_cells = table.rows[0].cells
    col_headers = ["Page", "Texte original", "Correction proposée", "Confiance", "Explication"]
    if is_h:
        col_headers.append("Source web")
    for i, label in enumerate(col_headers):
        hdr_cells[i].text = label
        _set_cell_bg(hdr_cells[i], "EEEEEE")
        para = hdr_cells[i].paragraphs[0]
        run = para.runs[0] if para.runs else para.add_run(label)
        run.font.name = _BODY_FONT
        run.font.size = _BODY_SIZE
        run.bold = True

    # Lignes de corrections (triées par page)
    sorted_corrections = sorted(cat_corrections, key=lambda c: getattr(c, "page_number", 0))

    confidence_colors = {
        "Certain":    "006600",
        "Probable":   "885500",
        "À vérifier": "AA0000",
    }

    for corr in sorted_corrections:
        page_number = getattr(corr, "page_number", 0)
        original_text = getattr(corr, "original_text", "") or ""
        corrected_text = getattr(corr, "corrected_text", "") or ""
        confidence = getattr(corr, "confidence", "Probable") or "Probable"
        explanation = getattr(corr, "explanation", "") or ""
        is_user_added = bool(getattr(corr, "is_user_added", False))

        # Tronquer l'explication pour la lisibilité du tableau
        if len(explanation) > _EXPLANATION_MAX_CHARS:
            explanation = explanation[:_EXPLANATION_MAX_CHARS].rstrip() + "…"

        # Nettoyer les marqueurs italique *...* pour le texte Word
        import re as _re
        original_text_clean = _re.sub(r'\*([^*]+)\*', r'\1', original_text[:300])
        corrected_text_clean = _re.sub(r'\*([^*]+)\*', r'\1', corrected_text[:300])

        row_cells = table.add_row().cells
        row_cells[0].text = str(page_number + 1)  # 1-indexed pour le lecteur
        row_cells[1].text = original_text_clean
        row_cells[2].text = corrected_text_clean
        row_cells[3].text = confidence
        row_cells[4].text = explanation
        if is_h:
            source_val = getattr(corr, "source", "") or ""
            row_cells[5].text = source_val[:200]

        for cell in row_cells:
            _set_cell_border(cell)
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.name = _BODY_FONT
                    run.font.size = _BODY_SIZE

        # Ligne éditeur : fond bleu très clair + note
        if is_user_added:
            for cell in row_cells:
                _set_cell_bg(cell, "EEF2FF")  # indigo-50 équivalent
            # Ajouter un indicateur dans la colonne explication
            expl_para = row_cells[4].paragraphs[0]
            note_text = " [Signalé par l'éditeur]"
            note_run = expl_para.add_run(note_text)
            note_run.italic = True
            note_run.font.name = _BODY_FONT
            note_run.font.size = Pt(9)
            note_run.font.color.rgb = RGBColor(0x43, 0x38, 0xCA)  # indigo

        # Couleur de la cellule confiance
        conf_hex = confidence_colors.get(confidence, "333333")
        for run in row_cells[3].paragraphs[0].runs:
            run.font.color.rgb = RGBColor(
                int(conf_hex[0:2], 16),
                int(conf_hex[2:4], 16),
                int(conf_hex[4:6], 16),
            )
            run.bold = True

        # Couleur de la cellule source (bleu URL si lien http, gris sinon)
        if is_h and row_cells[5].paragraphs[0].runs:
            src_run = row_cells[5].paragraphs[0].runs[0]
            src_text = src_run.text or ""
            if src_text.startswith("http"):
                src_run.font.color.rgb = RGBColor(0x00, 0x56, 0xB3)  # bleu lien
                src_run.underline = True
            elif src_text:
                src_run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    doc.add_paragraph()


# ── Point d'entrée public ──────────────────────────────────────────────────────


def export_corrections_docx(
    corrections: list,
    job: object,
    output_path: str,
) -> str:
    """
    Génère un rapport Word des corrections éditoriales pour un job donné.

    Paramètres
    ----------
    corrections :
        Liste d'objets Correction SQLAlchemy (modèle backend.models.Correction).
    job :
        Objet Job SQLAlchemy (modèle backend.models.Job).
    output_path :
        Chemin complet du fichier .docx à créer (répertoire parent doit exister).

    Retourne
    --------
    Chemin absolu du fichier généré (identique à output_path).

    Exemple
    -------
    >>> path = export_corrections_docx(corrections, job, "/tmp/rapport.docx")
    """
    # Calcul des statistiques par catégorie
    cat_counts: dict[str, int] = {}
    cat_map: dict[str, list] = {}
    for corr in corrections:
        cat = getattr(corr, "category", "H")
        cat_counts[cat] = cat_counts.get(cat, 0) + 1
        cat_map.setdefault(cat, []).append(corr)

    # Initialisation du document
    doc = Document()

    # Marges confortables
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.0)

    # ── 1. Page de titre ───────────────────────────────────────────────────────
    _add_title_page(doc, job, corrections, cat_counts)

    # ── 2. Tableau de synthèse ─────────────────────────────────────────────────
    doc.add_page_break()
    _add_summary_table(doc, cat_counts)

    # ── 3. Sections par catégorie ──────────────────────────────────────────────
    for cat in _CATEGORY_ORDER:
        if cat not in cat_map or not cat_map[cat]:
            continue
        doc.add_page_break()
        _add_category_section(doc, cat, cat_map[cat])

    # ── Sauvegarde ─────────────────────────────────────────────────────────────
    output_path_obj = Path(output_path)
    output_path_obj.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path_obj))

    logger.info(
        "Rapport DOCX généré : %s (%d corrections, %d catégorie(s))",
        output_path,
        len(corrections),
        len(cat_counts),
    )
    return str(output_path_obj.resolve())
